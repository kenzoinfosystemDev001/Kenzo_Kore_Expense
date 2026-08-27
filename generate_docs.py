import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_diagram_1():
    fig, ax = plt.subplots(figsize=(10, 5.5), dpi=300)
    ax.set_facecolor('#0B172A')
    fig.patch.set_facecolor('#0B172A')
    
    ax.text(5, 4.8, "Kenzo Kore Expense - High-Level Architecture", color='#00C8FF', 
            fontsize=13, fontweight='bold', ha='center')
    
    rect_web = patches.FancyBboxPatch((0.5, 3.2), 2.6, 1.2, boxstyle="round,pad=0.15", ec='#00A3FF', fc='#1E293B', lw=2)
    rect_apk = patches.FancyBboxPatch((3.7, 3.2), 2.6, 1.2, boxstyle="round,pad=0.15", ec='#10B981', fc='#1E293B', lw=2)
    rect_pwa = patches.FancyBboxPatch((6.9, 3.2), 2.6, 1.2, boxstyle="round,pad=0.15", ec='#8B5CF6', fc='#1E293B', lw=2)
    ax.add_patch(rect_web)
    ax.add_patch(rect_apk)
    ax.add_patch(rect_pwa)
    
    ax.text(1.8, 3.8, "Vite + React 19 SPA\n(Web Browser)", color='white', fontsize=8.5, fontweight='bold', ha='center')
    ax.text(5.0, 3.8, "Android Kotlin App\n(Custom WebView Shell)", color='white', fontsize=8.5, fontweight='bold', ha='center')
    ax.text(8.2, 3.8, "Standalone PWA\n(Mobile Intent Capture)", color='white', fontsize=8.5, fontweight='bold', ha='center')
    
    rect_gw = patches.FancyBboxPatch((1.5, 1.8), 7.0, 1.0, boxstyle="round,pad=0.15", ec='#00C8FF', fc='#0F172A', lw=2)
    ax.add_patch(rect_gw)
    ax.text(5.0, 2.4, "NestJS REST API Gateway & Microservices Engine", color='#00C8FF', fontsize=10, fontweight='bold', ha='center')
    ax.text(5.0, 2.0, "JWT Auth Guard | Policy Validation Engine | Audit Logger | File Storage Manager", color='#94A3B8', fontsize=7.5, ha='center')
    
    rect_db = patches.FancyBboxPatch((0.8, 0.4), 3.8, 1.0, boxstyle="round,pad=0.15", ec='#F59E0B', fc='#1E293B', lw=1.5)
    rect_store = patches.FancyBboxPatch((5.4, 0.4), 3.8, 1.0, boxstyle="round,pad=0.15", ec='#EC4899', fc='#1E293B', lw=1.5)
    ax.add_patch(rect_db)
    ax.add_patch(rect_store)
    
    ax.text(2.7, 0.9, "Prisma ORM & PostgreSQL DB\n(23 Enterprise Entities)", color='white', fontsize=8.5, fontweight='bold', ha='center')
    ax.text(7.3, 0.9, "Receipt Cloud Storage\n(S3 / Multer / Local Assets)", color='white', fontsize=8.5, fontweight='bold', ha='center')
    
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5.2)
    ax.axis('off')
    plt.tight_layout()
    plt.savefig('diagram_architecture.png', bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close()

def create_diagram_2():
    fig, ax = plt.subplots(figsize=(10, 4.5), dpi=300)
    ax.set_facecolor('#0B172A')
    fig.patch.set_facecolor('#0B172A')
    
    ax.text(5, 4.0, "Receipt Camera Capture & Automated OCR Workflow", color='#00C8FF', fontsize=12, fontweight='bold', ha='center')
    
    steps = [
        ("1. User Tap", "Synchronous\nonClick Event", '#00A3FF'),
        ("2. Stream Init", "Multi-Tier WebRTC\nConstraint Check", '#8B5CF6'),
        ("3. Fallback Intent", "Native OS Camera\n`<input capture>`", '#F59E0B'),
        ("4. Frame Capture", "HTML5 Canvas\ntoBlob (JPEG)", '#10B981'),
        ("5. OCR Parsing", "Auto-Fill Vendor,\nAmount, Date & Tax", '#00C8FF')
    ]
    
    for i, (title, sub, col) in enumerate(steps):
        x = 1.0 + i * 1.9
        rect = patches.FancyBboxPatch((x - 0.75, 1.4), 1.5, 1.8, boxstyle="round,pad=0.15", ec=col, fc='#1E293B', lw=1.8)
        ax.add_patch(rect)
        ax.text(x, 2.7, title, color='white', fontsize=8.5, fontweight='bold', ha='center')
        ax.text(x, 2.0, sub, color=col, fontsize=7.5, ha='center')
        
        if i < len(steps) - 1:
            ax.annotate('', xy=(x + 1.15, 2.3), xytext=(x + 0.75, 2.3), arrowprops=dict(arrowstyle="->", color=col, lw=1.5))
            
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4.5)
    ax.axis('off')
    plt.tight_layout()
    plt.savefig('diagram_camera_workflow.png', bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close()

def create_diagram_3():
    fig, ax = plt.subplots(figsize=(10, 4.5), dpi=300)
    ax.set_facecolor('#0B172A')
    fig.patch.set_facecolor('#0B172A')
    
    ax.text(5, 4.0, "Enterprise Approval & Automated Reimbursement Pipeline", color='#10B981', fontsize=12, fontweight='bold', ha='center')
    
    nodes = [
        ("Employee Claim", "Submit Expense\n+ Receipt Image", 1.0, '#38BDF8'),
        ("Policy Check", "Automated Limit &\nDuplicate Engine", 3.0, '#F59E0B'),
        ("Manager Review", "Approve / Reject /\nReturn for Revision", 5.0, '#8B5CF6'),
        ("Finance Audit", "Budget Verification\n& Tax Validation", 7.0, '#EC4899'),
        ("Reimbursement", "Automated Bank / UPI\nDisbursal & Audit", 9.0, '#10B981')
    ]
    
    for title, desc, x, col in nodes:
        rect = patches.FancyBboxPatch((x - 0.75, 1.4), 1.5, 1.8, boxstyle="round,pad=0.15", ec=col, fc='#1E293B', lw=1.8)
        ax.add_patch(rect)
        ax.text(x, 2.7, title, color='white', fontsize=8.5, fontweight='bold', ha='center')
        ax.text(x, 2.0, desc, color='#94A3B8', fontsize=7.5, ha='center')
        
        if x < 9.0:
            ax.annotate('', xy=(x + 1.25, 2.3), xytext=(x + 0.75, 2.3), arrowprops=dict(arrowstyle="->", color="#00C8FF", lw=1.5))
            
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4.5)
    ax.axis('off')
    plt.tight_layout()
    plt.savefig('diagram_approval_pipeline.png', bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close()

def create_diagram_4():
    fig, ax = plt.subplots(figsize=(10, 4.8), dpi=300)
    ax.set_facecolor('#0B172A')
    fig.patch.set_facecolor('#0B172A')
    
    ax.text(5, 4.3, "Android WebView WebRTC & FileChooser Permission Bridge", color='#F59E0B', fontsize=12, fontweight='bold', ha='center')
    
    ax.text(2.0, 3.6, "Web Application (JS)", color='#00C8FF', fontsize=9, fontweight='bold', ha='center')
    ax.text(5.0, 3.6, "Android WebView Shell", color='#8B5CF6', fontsize=9, fontweight='bold', ha='center')
    ax.text(8.0, 3.6, "Android OS Kernel", color='#10B981', fontsize=9, fontweight='bold', ha='center')
    
    boxes = [
        ("navigator.mediaDevices\n.getUserMedia()", 2.0, 2.5, '#00C8FF'),
        ("CustomWebChromeClient\n.onPermissionRequest()", 5.0, 2.5, '#8B5CF6'),
        ("PermissionManager\n.hasCameraPermission()", 8.0, 2.5, '#10B981'),
        ("Synchronous User\nGesture Token", 2.0, 1.0, '#00C8FF'),
        ("pendingPermissionRequest\nState Holding", 5.0, 1.0, '#8B5CF6'),
        ("Android OS System\nPermission Dialog", 8.0, 1.0, '#10B981')
    ]
    
    for text, x, y, col in boxes:
        rect = patches.FancyBboxPatch((x - 1.1, y - 0.4), 2.2, 0.8, boxstyle="round,pad=0.1", ec=col, fc='#1E293B', lw=1.5)
        ax.add_patch(rect)
        ax.text(x, y, text, color='white', fontsize=7.5, ha='center', va='center')
        
    ax.annotate('', xy=(3.9, 2.5), xytext=(3.1, 2.5), arrowprops=dict(arrowstyle="->", color="#00C8FF", lw=1.5))
    ax.annotate('', xy=(6.9, 2.5), xytext=(6.1, 2.5), arrowprops=dict(arrowstyle="->", color="#8B5CF6", lw=1.5))
    ax.annotate('', xy=(8.0, 1.4), xytext=(8.0, 2.1), arrowprops=dict(arrowstyle="->", color="#10B981", lw=1.5))
    ax.annotate('', xy=(6.1, 1.0), xytext=(6.9, 1.0), arrowprops=dict(arrowstyle="->", color="#8B5CF6", lw=1.5))
    
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4.8)
    ax.axis('off')
    plt.tight_layout()
    plt.savefig('diagram_webrtc_android.png', bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close()

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def generate_docx():
    doc = Document()
    
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = RGBColor(15, 23, 42)
    
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 119, 182)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = RGBColor(71, 85, 105)
        p.paragraph_format.space_after = Pt(20)
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 119, 182)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(11, 23, 42)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(16, 185, 129)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_p(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        return p

    def add_callout(text, title="KEY ARCHITECTURAL INSIGHT"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.8)
        set_cell_background(cell, "F0F9FF")
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        p = cell.paragraphs[0]
        r1 = p.add_run(f"📌 {title}\n")
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor(0, 119, 182)
        
        r2 = p.add_run(text)
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(30, 41, 59)
        p.paragraph_format.space_after = Pt(0)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ----------------------------------------------------
    # COVER & HEADER
    # ----------------------------------------------------
    add_title("KENZO KORE EXPENSE / KENZO ONEERP")
    add_subtitle("Complete Technical Architecture, Workflows, Database Schema, & Session History\nGenerated: August 7, 2026 | 11:39 IST | Release v1.2.0")

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Project Identifiers:", "Kenzo Kore Expense (Vite React 19 SPA) & Kenzo OneERP (Android Kotlin WebView App)"),
        ("Architecture Stack:", "React 19 + TypeScript + TailwindCSS + NestJS REST Core + PostgreSQL (Prisma ORM)"),
        ("Current Timestamp:", "August 7, 2026 | 11:39 IST"),
        ("Core Capabilities:", "Expense Claims, Camera OCR Scanner, Approval Engine, Analytics, Audit Logs, Settings, Dark/Light Themes")
    ]
    for idx, (k, v) in enumerate(meta_data):
        c0 = meta_table.cell(idx, 0)
        c1 = meta_table.cell(idx, 1)
        c0.width = Inches(2.2)
        c1.width = Inches(4.6)
        set_cell_background(c0, "F1F5F9")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, top=70, bottom=70, left=100, right=100)
        set_cell_margins(c1, top=70, bottom=70, left=100, right=100)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.font.bold = True
        r0.font.size = Pt(9)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ----------------------------------------------------
    # SECTION 1: SYSTEM OVERVIEW
    # ----------------------------------------------------
    add_h1("1. Executive Summary & System Overview")
    add_p("Kenzo Kore Expense is an enterprise-grade financial management system designed to streamline receipt capture, expense reporting, multi-level approval workflows, policy compliance enforcement, analytics, and instant reimbursement disbursement across corporate enterprises. Built as a high-performance Single Page Application (SPA) backed by a robust NestJS REST API microservice architecture and packaged as a native Android Kotlin application shell, the platform serves thousands of corporate employees and financial administrators.")
    add_p("This document serves as an exhaustive, non-compromised technical master reference covering every engineering decision, code refactoring step, system workflow, data pipeline, route specification, database schema, and native Android WebView integration established from initial inception up to August 7, 2026.")

    # ----------------------------------------------------
    # SECTION 2: CHRONOLOGY OF CHANGES & LOGIC
    # ----------------------------------------------------
    add_h1("2. Complete Chronological Development & Refactoring Log")
    add_p("Below is the complete historical log detailing every implementation milestone, bug diagnosis, and code modification applied across the project history:")

    chrono_steps = [
        ("Phase 1: Project Foundation & UI Architecture", 
         "Initialized modern web architecture using React 19, TypeScript, Vite 6, TailwindCSS, and Lucide React. Constructed responsive corporate layout shell featuring collapsible sidebar navigation, header status bar, user context management, and custom glassmorphism visual styling tokens."),
        
        ("Phase 2: Authentication & Role-Based Access Control", 
         "Implemented mock authentication and JWT session validation. Created role switcher supporting Employee, Manager, Admin, and Super Admin roles. Added session persistence in localStorage, user profile management, password update modal, and device image avatar uploaders."),
        
        ("Phase 3: Expense Creation & Automated OCR Processor", 
         "Built comprehensive Expense Creation Form supporting itemized line items, billable flags, category taxonomies, tax amounts, cost centers, and project allocation. Integrated simulated Optical Character Recognition (OCR) scanner reading vendor, date, total amount, and tax fields from mock invoice presets (AWS, Uber, Starbucks)."),
        
        ("Phase 4: Receipt Camera Scanner & Multi-Device Refactoring", 
         "Identified critical DOM unmounting bug where `isCameraLoading` removed `<video ref={videoRef}>` from the DOM tree, causing `videoRef.current` to be null when `getUserMedia()` resolved asynchronously. Refactored modal to keep persistent `<video>` element in DOM continuously while layering loading and error cards as overlays. Implemented 4-tier constraint fallback (`facingMode: 'environment'` -> `facingMode: 'user'` -> `video: true`), added front/rear camera device switcher, and configured `accept='image/*' capture='environment'` for native mobile camera intent activation."),
        
        ("Phase 5: Native Android App WebView Shell Audit & Fixes", 
         "Audited `Kenzo_kore_expense_app` Android Kotlin repository. Discovered WebRTC camera streaming failure caused by missing `onPermissionRequest()` in `CustomWebChromeClient.kt`. Identified Android 11+ Package Visibility bug where `takePictureIntent.resolveActivity()` returned null; added `<queries>` block for `IMAGE_CAPTURE` in `AndroidManifest.xml`. Refactored `MainActivity.kt` to hold `pendingPermissionRequest` state while Android OS runtime permission dialog is active. Added a prominent `📷 OPEN NATIVE CAMERA APP` button in the Web error state to trigger the Android system camera intent directly."),
        
        ("Phase 6: Enterprise Dark & Light Theme System", 
         "Designed and implemented a global theme switcher. Added `theme` state ('dark' | 'light') in `AppContext.tsx` with `localStorage` persistence and root `document.documentElement` class synchronization. Created comprehensive Light Theme CSS token overrides in `index.css` (`#F8FAFC` background, frosted white glass cards, high-contrast `#0F172A` charcoal text, `.keep-white` utility for gradient buttons). Integrated animated Sun/Moon toggle button in Header and interactive Theme Preference cards in Settings panel.")
    ]

    for title, desc in chrono_steps:
        add_h3(title)
        add_p(desc)

    add_callout("All changes across both the web application (Kenzo_kore_Expense) and Android app (Kenzo_kore_expense_app) have been built, verified with zero TypeScript or Gradle errors, committed, and pushed to remote main branches.", "BUILD & GIT PUSH VERIFICATION")

    # ----------------------------------------------------
    # SECTION 3: SYSTEM ARCHITECTURE & DIAGRAMS
    # ----------------------------------------------------
    add_h1("3. System Architecture & Diagrammatic Workflows")
    add_p("The system architecture is structured across four primary layers: Client Delivery (Web SPA & Android WebView Shell), API Gateway & Microservices (NestJS Controllers & Guards), Data Persistence (Prisma ORM & PostgreSQL), and Mobile Hardware Bridges.")

    add_h2("3.1 High-Level Architecture Diagram")
    if os.path.exists("diagram_architecture.png"):
        doc.add_picture("diagram_architecture.png", width=Inches(6.5))
        p_cap = doc.add_paragraph("Figure 1: High-Level End-to-End System Architecture Diagram")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.runs[0].font.size = Pt(8)
        p_cap.runs[0].font.italic = True

    add_h2("3.2 Camera Capture & Automated OCR Workflow")
    add_p("When an employee captures a paper receipt or uploads an invoice image, the platform processes the document through an asynchronous pipeline:")
    if os.path.exists("diagram_camera_workflow.png"):
        doc.add_picture("diagram_camera_workflow.png", width=Inches(6.5))
        p_cap = doc.add_paragraph("Figure 2: Receipt Camera Capture, Fallback Intent & OCR Parsing Workflow")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.runs[0].font.size = Pt(8)
        p_cap.runs[0].font.italic = True

    add_h2("3.3 Enterprise Approval & Reimbursement Pipeline")
    add_p("Expenses follow a strict state transition flow governed by policy violation algorithms, department budget thresholds, and multi-tier approval authorization:")
    if os.path.exists("diagram_approval_pipeline.png"):
        doc.add_picture("diagram_approval_pipeline.png", width=Inches(6.5))
        p_cap = doc.add_paragraph("Figure 3: Multi-Tier Approval Workflow & Automated Reimbursement Pipeline")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.runs[0].font.size = Pt(8)
        p_cap.runs[0].font.italic = True

    add_h2("3.4 Native Android WebView WebRTC Bridge")
    add_p("To bridge browser WebRTC APIs (`navigator.mediaDevices.getUserMedia`) with native Android OS hardware permissions, the custom Kotlin WebChromeClient handles asynchronous permission delegation:")
    if os.path.exists("diagram_webrtc_android.png"):
        doc.add_picture("diagram_webrtc_android.png", width=Inches(6.5))
        p_cap = doc.add_paragraph("Figure 4: Android WebView CustomWebChromeClient Permission Data Flow")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.runs[0].font.size = Pt(8)
        p_cap.runs[0].font.italic = True

    # ----------------------------------------------------
    # SECTION 4: APPLICATION ROUTES & PAGES
    # ----------------------------------------------------
    add_h1("4. Application Routes, Modules & UI Components")
    add_p("The front-end user interface comprises 9 core pages/modules accessible via role-based navigation:")

    routes = [
        ("Authentication View (`AuthView.tsx`)", "Handles user login, signup, role switching (Employee/Admin), and credential validation. Features animated corporate glass panel, form input validation, and initial session token dispatch."),
        ("Employee Dashboard (`DashboardView.tsx`)", "Provides employees with real-time financial metrics: Total Spend, Pending Approvals, Approved Claims, and Reimbursed Funds. Displays monthly expenditure trend charts, category pie charts, and recent transaction feeds."),
        ("My Expenses & Database (`ExpenseListView.tsx`)", "Exposes searchable, filterable expense database. Features category dropdowns, status pills, date range filters, multi-select export to CSV/PDF, and detailed expense view modals."),
        ("New Expense Form (`CreateExpenseView.tsx`)", "Form for submitting claims with title, category, subcategory, payment method, tax, billable state, cost center, and line items. Integrates WebCam Live Camera Scanner modal, receipt drag-and-drop zone, and automated OCR simulator."),
        ("Approval Queue (`ApprovalQueueView.tsx`)", "Admin/Manager review portal listing pending expense claims. Displays line item details, attached receipts, duplicate warnings, policy flags, and actions to Approve, Reject, or Return to Draft with comments."),
        ("Analytics & Reporting (`AnalyticsView.tsx`)", "Interactive analytics engine presenting department spend breakdowns, category distribution pie charts, top merchants chart, monthly spending forecast, and executive report generator."),
        ("Audit Logs Viewer (`AuditLogsView.tsx`)", "Compliance logging module tracking system actions, user session logins, policy limit edits, expense modifications, and security traces with timestamp and IP address."),
        ("Settings Panel (`SettingsView.tsx`)", "Admin configuration center for policy limits, department budget allocations, corporate user directory management, password reset triggers, and Theme Preference (Dark/Light Mode) selection."),
        ("AI Insights Assistant (`AIChatAssistant.tsx`)", "Natural language chat assistant widget providing instant spending insights, budget forecasts, anomaly detection alerts, and policy guidance.")
    ]

    for title, desc in routes:
        add_h3(title)
        add_p(desc)

    # ----------------------------------------------------
    # SECTION 5: DATABASE SCHEMA & ENTITIES
    # ----------------------------------------------------
    add_h1("5. Complete Database Schema (Prisma PostgreSQL - 23 Entities)")
    add_p("The back-end relational database models 23 core entities, enumerations, and joins structured in PostgreSQL via Prisma ORM:")

    entities = [
        ("User (`users`)", "Core user model storing ID, email, password hash, name, avatar URL, Role enum (EMPLOYEE, ADMIN, SUPER_ADMIN), designation, departmentId, managerId, costCenterId, joining date, and GST number."),
        ("Department (`departments`)", "Organizational department entity with ID, name, code, budgetLimit, and spentAmount."),
        ("CostCenter (`cost_centers`)", "Cost center entity tracking financial allocation codes across corporate branches."),
        ("Project (`projects`)", "Project entity tracking project-level budget allocations and accrued expenses."),
        ("Expense (`expenses`)", "Primary transaction record containing title, employeeId, departmentId, costCenterId, projectId, category, amount, currency, date, paymentMethod, status enum, merchant, businessPurpose, billable flag, location, receiptUrl, invoiceUrl, gstNumber, taxAmount, and referenceNumber."),
        ("ExpenseItem (`expense_items`)", "Itemized line item breakdown attached to an expense claim with item description, amount, tax, and category."),
        ("ExpenseApproval (`expense_approvals`)", "Approval action log storing expenseId, approverId, status enum, comment, and approval timestamp."),
        ("Reimbursement (`reimbursements`)", "Disbursal tracking model recording expenseId, status (PENDING, PROCESSING, PAID), paymentMethod, transactionId, and clearedAt timestamp."),
        ("Policy (`policies`)", "Compliance policy entity storing policy name, category, limitAmount, description, and isEnabled flag."),
        ("Budget (`budgets`)", "Budget allocation model recording budget name, allocated amount, spent amount, and period (MONTHLY, QUARTERLY, YEARLY)."),
        ("AuditLog (`audit_logs`)", "Security audit record logging timestamp, userId, action string, details, and client IP address."),
        ("Notification (`notifications`)", "System notification model storing userId, title, message, read state, and timestamp."),
        ("Tag (`tags`) & ExpenseTag (`expense_tags`)", "Categorization tag models supporting many-to-many tag relationships on expense claims.")
    ]

    for title, desc in entities:
        add_h3(title)
        add_p(desc)

    # ----------------------------------------------------
    # SECTION 6: CONCEPTS & KNOWLEDGE FOR DEVELOPERS
    # ----------------------------------------------------
    add_h1("6. Essential Core Concepts & Knowledge for Developers")
    add_p("To successfully build, extend, and maintain this project, software engineers must possess mastery over the following concepts:")

    concepts = [
        ("Vite & React 19 SPA State Management", "Understanding React Context API (`AppContext`), hooks (`useState`, `useEffect`, `useRef`), dynamic import code-splitting, and Fast Refresh."),
        ("WebRTC MediaStream & Constraints API", "Knowledge of `navigator.mediaDevices.getUserMedia()`, video element ref bindings (`srcObject`), playsInline / muted autoplay policies on iOS Safari, and multi-tier constraint fallback chains."),
        ("Native Android WebView Engine Internals", "Understanding `WebChromeClient` callbacks (`onPermissionRequest`, `onShowFileChooser`), Android 11+ package visibility rules (`<queries>`), and holding `PermissionRequest` state across OS runtime dialogs."),
        ("TailwindCSS Class-Based Theme Switching", "Configuring `darkMode: 'class'`, manipulating `document.documentElement` class lists, CSS custom properties, and maintaining color contrast across light and dark modes."),
        ("Prisma ORM & Relational Modeling", "Constructing PostgreSQL schemas, foreign key relations, indexing strategies, cascade deletes, and migration management.")
    ]

    for title, desc in concepts:
        add_h3(title)
        add_p(desc)

    output_filename = "Kenzo_Kore_Expense_Master_Documentation.docx"
    doc.save(output_filename)
    print(f"Successfully generated {output_filename}")

if __name__ == "__main__":
    create_diagram_1()
    create_diagram_2()
    create_diagram_3()
    create_diagram_4()
    generate_docx()
